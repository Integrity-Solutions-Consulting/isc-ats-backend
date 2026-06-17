"""Word profile document generator.

Given an application ID, builds a filled .docx using Integrity Solutions'
candidate profile template. Uses `candidate.parsed_data` as the CV data
source; triggers a fresh CV parse if parsed_data is absent.

Sections filled:
  1. Personal data   — from recruitment.candidates + auth.users email
  2. Work experience — from parsed_data.experience (dynamic table count)
  3. Technical knowledge — from parsed_data.skills + parsed_data.tools (bullets)
  4. Skills          — from parsed_data.soft_skills (bullets)
  5. Certifications  — from parsed_data.certifications
  6. Projects        — from parsed_data.projects
"""

from __future__ import annotations

import copy
import io
import logging
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import app.models_registry  # noqa: F401
from app.core.database import async_session_factory
from app.modules.auth.infrastructure.models import User
from app.modules.org.infrastructure.models import Parameter
from app.modules.recruitment.infrastructure.application_models import Application
from app.modules.recruitment.infrastructure.candidate_models import Candidate
from app.modules.recruitment.infrastructure.models import Vacancy
from app.modules.ai.application.cv_parse_service import parse_candidate_cv

logger = logging.getLogger(__name__)

TEMPLATE_PATH = Path(__file__).parent.parent.parent.parent / "templates" / "perfil_candidato.docx"

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


# ─── XML helpers ─────────────────────────────────────────────────────────────

def _set_cell_text(cell: Any, text: str) -> None:
    """Set cell text, preserving the paragraph but replacing all runs."""
    para = cell.paragraphs[0] if cell.paragraphs else None
    if para is None:
        return
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = text or ""
    elif text:
        para.add_run(text)


def _fill_simple_table(table: Any, values: dict[int, str]) -> None:
    """Fill row_index → value into the second column of a 2-col table."""
    for row_idx, value in values.items():
        if row_idx < len(table.rows):
            _set_cell_text(table.rows[row_idx].cells[1], value)


def _fill_exp_table_elem(table_elem: Any, exp: dict[str, Any]) -> None:
    """Fill an experience table XML element (works on raw lxml element)."""
    rows = table_elem.findall(f"{{{W}}}tr")
    if not rows:
        return

    def set_row(idx: int, text: str) -> None:
        if idx >= len(rows):
            return
        cells = rows[idx].findall(f"{{{W}}}tc")
        if len(cells) < 2:
            return
        cell = cells[1]
        p = cell.find(f"{{{W}}}p")
        if p is None:
            return
        for r in p.findall(f"{{{W}}}r"):
            p.remove(r)
        if text:
            r_elem = OxmlElement("w:r")
            t_elem = OxmlElement("w:t")
            if text.startswith(" ") or text.endswith(" ") or "\n" in text:
                t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            t_elem.text = text
            r_elem.append(t_elem)
            p.append(r_elem)

    set_row(0, exp.get("company", ""))
    set_row(1, exp.get("role", ""))
    set_row(2, exp.get("start_date", ""))
    set_row(3, exp.get("end_date", ""))
    functions = exp.get("functions", [])
    set_row(4, "\n".join(f"• {f}" for f in functions) if functions else "")
    tools = exp.get("tools", [])
    set_row(5, ", ".join(tools) if tools else "")


def _fill_experience_section(doc: Document, experiences: list[dict]) -> None:
    """Manage experience tables: fill N, remove extras, add more if needed."""
    body = doc.element.body
    n = len(experiences)

    # Tables 1, 2, 3 are the pre-built experience blocks.
    pre_built = doc.tables[1:4]
    template_copy = copy.deepcopy(pre_built[0]._element)

    if n == 0:
        for t in pre_built:
            body.remove(t._element)
        return

    # Fill first experience into pre-built table[1]
    _fill_exp_table_elem(pre_built[0]._element, experiences[0])

    # Remove tables[2] and tables[3] — we'll add back dynamically
    for t in pre_built[1:]:
        body.remove(t._element)

    # Append additional experience tables after the first one
    anchor = pre_built[0]._element
    for exp in experiences[1:]:
        new_elem = copy.deepcopy(template_copy)
        _fill_exp_table_elem(new_elem, exp)
        anchor.addnext(new_elem)
        anchor = new_elem


def _find_bullet_placeholder(doc: Document, heading_marker: str) -> Any | None:
    """Locate the first List Paragraph after the given section heading.

    Paragraph indexes shift whenever the template (or an earlier section)
    changes, so placeholders are found by scanning, never by position.
    """
    heading_found = False
    for para in doc.paragraphs:
        if heading_marker in para.text:
            heading_found = True
            continue
        if heading_found and para.style.name == "List Paragraph":
            return para
    return None


def _fill_bullet_section(doc: Document, placeholder_para: Any, items: list[str]) -> None:
    """Replace a List Paragraph placeholder with actual bullet items."""
    body = doc.element.body
    placeholder_elem = placeholder_para._element

    if not items:
        body.remove(placeholder_elem)
        return

    # Insert bullet paragraphs before the placeholder, then remove placeholder
    for item in reversed(items):
        new_p = copy.deepcopy(placeholder_elem)
        # Clear runs and set text
        for r in new_p.findall(f"{{{W}}}r"):
            new_p.remove(r)
        r_elem = OxmlElement("w:r")
        t_elem = OxmlElement("w:t")
        t_elem.text = item
        r_elem.append(t_elem)
        new_p.append(r_elem)
        placeholder_elem.addprevious(new_p)

    body.remove(placeholder_elem)


def _fill_certifications_table(doc: Document, certifications: list[dict]) -> None:
    """Fill the certifications table (index 4 after experience tables adjusted)."""
    # After adjusting experience tables, re-resolve by scanning for known header
    cert_table = None
    for table in doc.tables:
        if table.rows and table.rows[0].cells[0].text.strip() == "Institución":
            cert_table = table
            break
    if cert_table is None:
        return

    # The table has a header row + 3 data rows by default.
    # We need exactly len(certifications) data rows.
    header_row = cert_table.rows[0]._element

    # Remove all existing data rows
    for row in cert_table.rows[1:]:
        cert_table._element.remove(row._element)

    if not certifications:
        # Add one empty row so the table isn't just a header
        _add_cert_row(cert_table, {})
        return

    for cert in certifications:
        _add_cert_row(cert_table, cert)


def _add_cert_row(table: Any, cert: dict) -> None:
    from docx.oxml import OxmlElement
    # Copy structure from header row to create a data row
    header_row_elem = table.rows[0]._element
    new_row = copy.deepcopy(header_row_elem)
    # Clear header style markers if any and set text
    cells = new_row.findall(f"{{{W}}}tc")
    values = [
        cert.get("institution", ""),
        cert.get("name", ""),
        cert.get("start", ""),
        cert.get("end", ""),
    ]
    for i, cell in enumerate(cells):
        p = cell.find(f"{{{W}}}p")
        if p is not None:
            for r in p.findall(f"{{{W}}}r"):
                p.remove(r)
            if i < len(values) and values[i]:
                r_elem = OxmlElement("w:r")
                t_elem = OxmlElement("w:t")
                t_elem.text = values[i]
                r_elem.append(t_elem)
                p.append(r_elem)
    table._element.append(new_row)


def _fill_projects_table(doc: Document, projects: list[dict]) -> None:
    """Fill the projects table (identified by 'Tema' header)."""
    proj_table = None
    for table in doc.tables:
        if table.rows and table.rows[0].cells[0].text.strip() == "Tema":
            proj_table = table
            break
    if proj_table is None:
        return

    for row in proj_table.rows[1:]:
        proj_table._element.remove(row._element)

    if not projects:
        _add_proj_row(proj_table, {})
        return

    for proj in projects:
        _add_proj_row(proj_table, proj)


def _add_proj_row(table: Any, proj: dict) -> None:
    header_row_elem = table.rows[0]._element
    new_row = copy.deepcopy(header_row_elem)
    cells = new_row.findall(f"{{{W}}}tc")
    tools = proj.get("tools", [])
    values = [
        proj.get("name", ""),
        proj.get("description", ""),
        ", ".join(tools) if tools else "",
    ]
    for i, cell in enumerate(cells):
        p = cell.find(f"{{{W}}}p")
        if p is not None:
            for r in p.findall(f"{{{W}}}r"):
                p.remove(r)
            if i < len(values) and values[i]:
                r_elem = OxmlElement("w:r")
                t_elem = OxmlElement("w:t")
                t_elem.text = values[i]
                r_elem.append(t_elem)
                p.append(r_elem)
    table._element.append(new_row)


# ─── Main entry point ────────────────────────────────────────────────────────

async def generate_profile_word(application_id: int) -> bytes:
    """Build and return the filled .docx as bytes."""
    async with async_session_factory() as session:
        app: Application | None = await session.get(Application, application_id)
        if not app:
            raise ValueError(f"Application {application_id} not found")

        candidate: Candidate | None = await session.get(Candidate, app.candidate_id)
        if not candidate:
            raise ValueError(f"Candidate for application {application_id} not found")

        vacancy: Vacancy | None = await session.get(Vacancy, app.vacancy_id)
        user: User | None = await session.get(User, candidate.user_id)

        # Resolve city/province names
        city_name = ""
        province_name = ""
        if candidate.city_id:
            city_param: Parameter | None = await session.get(Parameter, candidate.city_id)
            city_name = city_param.name if city_param else ""
        if candidate.province_id:
            prov_param: Parameter | None = await session.get(Parameter, candidate.province_id)
            province_name = prov_param.name if prov_param else ""

        # Resolve university name
        university_name = ""
        if candidate.university_id:
            univ_param: Parameter | None = await session.get(Parameter, candidate.university_id)
            university_name = univ_param.name if univ_param else ""

        # Resolve vacancy position name
        position_name = ""
        if vacancy and vacancy.vacancy_name_id:
            vac_param: Parameter | None = await session.get(Parameter, vacancy.vacancy_name_id)
            position_name = vac_param.name if vac_param else ""

    # Trigger CV parse if no parsed_data yet
    parsed: dict[str, Any] | None = candidate.parsed_data
    if not parsed:
        logger.info("No parsed_data for candidate %d — triggering parse", candidate.id)
        parsed = await parse_candidate_cv(candidate.id)
        if not parsed:
            parsed = {}

    # ─── Build document ──────────────────────────────────────────────────────
    doc = Document(TEMPLATE_PATH)

    full_name = f"{candidate.first_name} {candidate.last_name}".strip()
    location = f"{city_name}, {province_name}".strip(", ") if (city_name or province_name) else ""

    # Section 1 — personal data (Table 0)
    _fill_simple_table(doc.tables[0], {
        0: full_name,
        1: position_name,
        2: candidate.degree_title or "",     # Título Obtenido
        3: university_name,                 # Universidad
        4: location,
        5: user.email if user else "",
        6: candidate.cedula or "",
        7: candidate.phone or "",
        8: candidate.home_address or "",    # Dirección Domiciliaria
        9: "",                              # Dirección GPS — not in DB
    })

    # Section 2 — experience
    experiences: list[dict] = parsed.get("experience", [])
    _fill_experience_section(doc, experiences)

    # Section 3 — technical knowledge
    skills: list[str] = parsed.get("skills", [])
    tools: list[str] = parsed.get("tools", [])
    tech_items = skills + [t for t in tools if t not in skills]
    tech_placeholder = _find_bullet_placeholder(doc, "3. CONOCIMIENTOS")
    if tech_placeholder is not None:
        _fill_bullet_section(doc, tech_placeholder, tech_items)

    # Section 4 — soft skills
    soft_skills: list[str] = parsed.get("soft_skills", [])
    habilidades_placeholder = _find_bullet_placeholder(doc, "4. HABILIDADES")
    if habilidades_placeholder is not None:
        _fill_bullet_section(doc, habilidades_placeholder, soft_skills)

    # Section 5 — certifications
    certifications: list[dict] = parsed.get("certifications", [])
    _fill_certifications_table(doc, certifications)

    # Section 6 — projects
    projects: list[dict] = parsed.get("projects", [])
    _fill_projects_table(doc, projects)

    # ─── Serialize ───────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
